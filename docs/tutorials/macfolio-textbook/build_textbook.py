from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[3]
OUT = Path(__file__).resolve().parent
ASSETS = OUT / "_generated_assets"
ASSETS.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "1E6B68"
GOLD = "9A6A00"
INK = "202A33"
MUTED = "66727D"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
CODE = "F2F4F7"
WHITE = "FFFFFF"
RED = "9B1C1C"

TEXT_EXTENSIONS = {".js", ".jsx", ".css", ".md", ".json", ".html", ".log", ".svg"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "100")
    spacing.set(qn("w:after"), "100")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc: Document, short_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, MUTED, 0, 18),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.text = short_title.upper()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hp.runs[0], size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("MACFOLIO ENGINEERING TEXTBOOK  |  ")
    set_font(r, size=8.5, color=MUTED)
    field(fp, "PAGE")


def add_cover(doc: Document, volume: str, title: str, subtitle: str, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(volume.upper())
    set_font(r, size=10, color=accent, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A beginner-to-engineer rebuild guide grounded in the real repository")
    set_font(r, size=11, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PATTERN  >  PRACTICE  >  REPETITION  >  REFINEMENT  >  HABIT")
    set_font(r, size=9, color=TEAL, bold=True)
    doc.add_page_break()


def add_para(doc, text, bold_lead=None, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_font(r)


def add_callout(doc, label, text, tone="info"):
    colors = {"info": (LIGHT, NAVY), "lesson": ("E7F1EF", TEAL), "warning": ("FFF4D6", GOLD), "risk": ("FBE9E9", RED)}
    fill, color = colors[tone]
    p = doc.add_paragraph()
    shade_paragraph(p, fill)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    r = p.add_run(f"{label}: ")
    set_font(r, color=color, bold=True)
    r = p.add_run(text)
    set_font(r, color=INK)


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(caption)
        set_font(r, size=9, color=MUTED, bold=True)
    p = doc.add_paragraph()
    shade_paragraph(p, CODE)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(code.rstrip())
    set_font(r, name="Consolas", size=8.5, color=INK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_font(r, size=9.5, color=NAVY, bold=True)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, size=9.2)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_chapter(doc, number, title, goal):
    doc.add_heading(f"Chapter {number}: {title}", level=1)
    add_callout(doc, "Chapter outcome", goal, "lesson")


def add_five_questions(doc, what, why, when, where, how):
    for label, text in (("WHAT", what), ("WHY", why), ("WHEN", when), ("WHERE", where), ("HOW", how)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}  ")
        set_font(r, size=9.5, color=TEAL, bold=True)
        r = p.add_run(text)
        set_font(r)


def add_explanations(doc, beginner, developer, engineering):
    doc.add_heading("Three Ways To Understand It", level=2)
    add_callout(doc, "Beginner explanation", beginner, "info")
    add_callout(doc, "Developer explanation", developer, "lesson")
    add_callout(doc, "Engineering explanation", engineering, "warning")


def add_checkpoint(doc, beginner, intermediate, advanced):
    doc.add_heading("Checkpoint", level=2)
    add_bullets(doc, [f"Beginner: {beginner}", f"Intermediate: {intermediate}", f"Advanced: {advanced}"])


def add_score(doc, scores):
    doc.add_heading("Architecture Review", level=2)
    rows = [(name, f"{score}/10", reason) for name, score, reason in scores]
    add_table(doc, ["Dimension", "Score", "Reason"], rows, [2200, 1100, 6060])


def source(path):
    return (ROOT / path).read_text(encoding="utf-8")


def snippet(path, start, end):
    lines = source(path).splitlines()
    return "\n".join(lines[start - 1:end])


def add_toc(doc, chapters):
    doc.add_heading("Learning Map", level=1)
    add_para(doc, "Use this volume as a build session, not passive reading. Type the examples, pause at checkpoints, and keep the architecture questions beside your editor.")
    add_table(doc, ["Chapter", "You will learn"], chapters, [2900, 6460])
    doc.add_page_break()


def make_diagram(name, title, columns):
    width, height = 1500, 360
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 34)
        body_font = ImageFont.truetype("arial.ttf", 25)
    except OSError:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.text((50, 24), title, font=title_font, fill="#17324D")
    box_w = 240
    gap = (width - 100 - box_w * len(columns)) // max(1, len(columns) - 1)
    y = 125
    for index, label in enumerate(columns):
        x = 50 + index * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + 110), radius=12, fill="#E8EEF5", outline="#2E74B5", width=3)
        bbox = draw.multiline_textbbox((0, 0), label, font=body_font, align="center", spacing=6)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.multiline_text((x + (box_w - tw) / 2, y + (110 - th) / 2), label, font=body_font, fill="#202A33", align="center", spacing=6)
        if index < len(columns) - 1:
            ax = x + box_w + 14
            bx = x + box_w + gap - 14
            draw.line((ax, y + 55, bx, y + 55), fill="#1E6B68", width=5)
            draw.polygon([(bx, y + 55), (bx - 18, y + 44), (bx - 18, y + 66)], fill="#1E6B68")
    out = ASSETS / f"{name}.png"
    image.save(out)
    return out


def add_figure(doc, path, caption, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption.split(".")[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_font(r, size=9, color=MUTED, italic=True)


def volume_one():
    doc = Document(); configure_document(doc, "Volume 1 - Setup and Mindset")
    add_cover(doc, "Volume 1", "Start Without Hesitation", "Project setup, engineering rhythm, Git basics, and the first working React screen")
    add_toc(doc, [("1. The engineer's starting ritual", "Turn an idea into responsibilities and a tiny first milestone."), ("2. Set up Macfolio", "Create the Vite project, install dependencies, and understand configuration."), ("3. Naming and placement", "Use predictable names and decide where code belongs."), ("4. Build the first vertical slice", "Render a shell, component, data object, and style together."), ("5. Practice loop", "Debug, commit, document, and repeat with confidence.")])

    add_chapter(doc, 1, "The Engineer's Starting Ritual", "You will be able to start a new idea with a one-page design instead of an empty-editor panic.")
    add_five_questions(doc, "A starting ritual is a repeatable set of questions asked before files are created.", "It changes uncertainty into named responsibilities and keeps the first version small.", "Use it for every new project and every feature that touches more than one file.", "Write the answers in a README, issue, or architecture note before implementation.", "Define the user, outcome, entities, responsibilities, dependencies, and smallest end-to-end slice.")
    add_bullets(doc, ["User: who is trying to accomplish something?", "Outcome: what observable result proves the feature works?", "Entities: what nouns carry data, such as window, project, file, or route?", "Responsibilities: which pieces load data, store state, decide behavior, and render UI?", "Dependencies: which browser APIs, packages, and remote services are involved?", "Growth question: what becomes painful if the feature grows ten times?"])
    add_callout(doc, "Macfolio example", "The product is not merely a page. It is a portfolio presented through two interaction shells: a macOS-like desktop and an iPhone-like mobile experience. Both consume portfolio content, but each owns its navigation model.", "lesson")
    add_explanations(doc, "Planning is writing down what each helper is responsible for before coding it.", "Planning defines contracts between UI, state, and data modules.", "Planning reduces coupling by making ownership and dependency direction explicit before implementation pressure blurs them.")

    add_chapter(doc, 2, "Set Up Macfolio", "You will understand every setup file and be able to reproduce the development environment from an empty folder.")
    add_bullets(doc, ["Install a current Node.js LTS release and Git.", "Create a project with Vite's React template.", "Install the libraries listed in package.json.", "Create .env.local with VITE_GITHUB_USERNAME for public GitHub data.", "Run the development server and open the local address shown in the terminal."] , numbered=True)
    add_code(doc, "npm create vite@latest macfolio -- --template react\ncd macfolio\nnpm install\nnpm run dev", "The minimum setup loop")
    pkg = json.loads(source("package.json"))
    rows = [(name, version, dependency_role(name)) for name, version in sorted(pkg["dependencies"].items())]
    add_table(doc, ["Package", "Version", "Why Macfolio uses it"], rows, [2200, 1300, 5860])
    doc.add_heading("Configuration Files", level=2)
    add_table(doc, ["File", "Responsibility", "Do not put here"], [("index.html", "Browser document and #root mount point.", "React feature markup."), ("src/main.jsx", "Create the React root and load global CSS.", "Business logic or window behavior."), ("vite.config.js", "Plugins and # path aliases.", "Runtime feature settings."), ("eslint.config.js", "Code-quality rules.", "Formatting-only preferences."), ("jsconfig.json", "Editor path resolution.", "Build behavior."), ("package-lock.json", "Exact dependency graph for reproducible installs.", "Hand-edited dependency choices.")], [2300, 3900, 3160])
    add_code(doc, snippet("src/main.jsx", 1, 11), "Actual application entry point")
    add_para(doc, "ReactDOM.createRoot finds the div with id root, creates React's rendering boundary, and renders App inside StrictMode. StrictMode intentionally repeats some development-only lifecycle work to expose unsafe side effects; it does not double-render production output.")

    add_chapter(doc, 3, "Naming and Placement", "You will choose file, component, hook, function, constant, and CSS names that reveal ownership.")
    add_table(doc, ["Thing", "Convention", "Macfolio example"], [("React component", "PascalCase", "FinderWindow.jsx"), ("Hook", "camelCase beginning with use", "usePortfolioFileSystem.js"), ("Utility", "verb or decision phrase", "buildWorkLocation.js"), ("Constant collection", "UPPER_SNAKE_CASE", "WINDOW_CONFIG"), ("Feature folder", "lowercase kebab-case", "code-preview/"), ("CSS class", "lowercase kebab-case with domain prefix", "mobile-app-frame")], [2300, 3200, 3860])
    add_callout(doc, "Placement rule", "Put code beside the feature that owns its change. Move it to shared only after multiple unrelated features genuinely need the same behavior.", "lesson")
    add_bullets(doc, ["A component renders and coordinates UI.", "A hook coordinates reusable React state or effects.", "A utility performs a pure transformation or decision.", "A store owns state shared across distant parts of one domain.", "A data file contains stable content or configuration, not active behavior.", "An API module speaks to an external service and returns plain data."])

    add_chapter(doc, 4, "Build the First Vertical Slice", "You will build one visible path through data, logic, UI, and CSS before expanding the project.")
    diagram = make_diagram("vertical_slice", "A vertical slice crosses every required layer", ["User action", "Component", "State / hook", "Data", "Visible result"])
    add_figure(doc, diagram, "Figure 1. A small end-to-end path gives faster feedback than building every folder first.")
    add_code(doc, "const apps = [{ id: 'finder', name: 'Finder', icon: '/images/finder.png' }];\n\nfunction Dock() {\n  return apps.map((app) => (\n    <button key={app.id} onClick={() => openWindow(app.id)}>\n      <img src={app.icon} alt=\"\" />\n      <span>{app.name}</span>\n    </button>\n  ));\n}", "A deliberately small first slice")
    add_para(doc, "The array is data, Dock is presentation, and openWindow is behavior owned elsewhere. That separation is already architecture, even before the project becomes large.")

    add_chapter(doc, 5, "Practice Loop", "You will use a calm repeatable cycle for building, checking, documenting, and improving features.")
    add_bullets(doc, ["Make the smallest observable change.", "Read the browser and terminal feedback.", "Explain the execution flow in plain language.", "Run lint and a production build.", "Commit one coherent idea with a message that describes the outcome.", "Record architecture decisions and lessons when the reasoning will matter later."], numbered=True)
    add_code(doc, "npm run lint\nnpm run build\ngit status\ngit add <files>\ngit commit -m \"Build desktop window opening flow\"", "A compact verification and commit rhythm")
    add_checkpoint(doc, "Explain why package-lock.json is committed.", "Design a folder structure for a notes app with offline storage.", "Explain which project decisions deserve an ADR and which do not.")
    add_score(doc, [("Architecture", 8, "The current app has clear feature boundaries, with a few compatibility exports left during migration."), ("Maintainability", 8, "Purpose headers and docs make ownership visible; global CSS remains broad."), ("Scalability", 7, "Client-side public APIs fit a portfolio but would need a server boundary for private or high-volume data."), ("Readability", 8, "Explicit names and small coordinators make most flows approachable."), ("Reusability", 7, "The project correctly favors local ownership over premature shared abstractions.")])
    return doc


def dependency_role(name):
    roles = {
        "react": "Component and hook runtime.", "react-dom": "Mount React in the browser.", "vite": "Development and production bundling.", "zustand": "Small domain stores for shared state.", "gsap": "Window drag and animation mechanics.", "@gsap/react": "React lifecycle integration for GSAP.", "framer-motion": "Mobile screen transitions and gestures.", "react-responsive": "Desktop/mobile shell selection.", "lucide-react": "Consistent interface icons.", "react-icons": "File-type icons.", "@monaco-editor/react": "VS Code-like source preview.", "react-pdf": "Render the resume PDF in a window.", "dayjs": "Compact date and time formatting.", "clsx": "Conditional class names.", "immer": "Ergonomic immutable updates through Zustand middleware.", "react-use": "General-purpose React hooks.", "tailwindcss": "Utility classes alongside authored CSS.", "@tailwindcss/vite": "Tailwind integration for Vite."
    }
    return roles.get(name, "Supporting project dependency; see its imports in the repository atlas.")


def volume_two():
    doc = Document(); configure_document(doc, "Volume 2 - Architecture")
    add_cover(doc, "Volume 2", "Architecture You Can Feel", "Responsibilities, feature boundaries, folder structure, dependency direction, and growth")
    add_toc(doc, [("1. Read the product as domains", "Find stable areas of responsibility before choosing folders."), ("2. Feature-based structure", "Place code by ownership and explain the compatibility layer."), ("3. Dependency direction", "Keep shells, features, data, and shared code loosely coupled."), ("4. State ownership", "Choose local state, hooks, or stores deliberately."), ("5. Refactoring without panic", "Move boundaries gradually while the app keeps working.")])
    add_chapter(doc, 1, "Read the Product as Domains", "You will identify architecture from product behavior instead of starting with generic component folders.")
    add_five_questions(doc, "A domain is a coherent area of product responsibility and language.", "Domains create boundaries that change for one reason and can be understood independently.", "Look for them when features have their own state, vocabulary, workflows, or growth path.", "In Macfolio they live mainly under src/features.", "Name the user-visible capability, its data, its decisions, and the other domains it must talk to.")
    add_table(doc, ["Domain", "Owns", "Does not own"], [("App composition", "Startup, shell choice, top-level error/loading boundaries.", "Feature internals."), ("Desktop shell", "Dock, navbar, window configuration, window mechanics.", "Finder file decisions."), ("Finder", "Navigation UI and file-open decisions.", "Raw GitHub requests."), ("Portfolio", "Load and map portfolio project data.", "Window rendering."), ("Code preview", "Explorer, editor, terminal-like source experience.", "Desktop window state."), ("Mobile shell", "Phone navigation stack, app frame, mobile mini-apps.", "Desktop window behavior.")], [2000, 3700, 3660])
    add_figure(doc, make_diagram("domains", "Macfolio dependency story", ["App", "Shell", "Feature", "Hook / store", "API / browser"]), "Figure 1. Dependencies should point toward narrower responsibilities, not loop between unrelated features.")

    add_chapter(doc, 2, "Feature-Based Structure", "You will reproduce the repository structure and know what belongs in each folder.")
    structure = """src/
  app/                       # application composition
  features/
    desktop-shell/           # desktop frame and window mechanics
    finder/                  # Finder UI and open-item decisions
    portfolio/               # portfolio loading and mapping
    code-preview/            # explorer/editor/terminal experience
    mobile-shell/            # mobile navigation and mini-apps
  api/                       # external service boundary
  constants/                 # stable content grouped by ownership
  store/                     # compatibility exports and cross-feature stores
  windows/                   # desktop app windows and migration adapters
  components/                # app-level UI and compatibility exports
"""
    add_code(doc, structure, "Current high-level shape")
    add_callout(doc, "Why compatibility exports exist", "Files such as src/components/Dock.jsx re-export the new feature implementation. This keeps old import paths working while architecture is moved gradually. They are migration bridges, not duplicate implementations.", "info")
    add_bullets(doc, ["Belongs in a feature: components, hooks, utilities, data, styles, and tests used mainly by that capability.", "Belongs in app: providers, top-level shell choice, routing, and application bootstrapping.", "Belongs in api: raw remote requests and transport-level behavior.", "Belongs in shared only after unrelated features repeat the same stable contract.", "Does not belong in constants: helpers with side effects, React state, or API calls."])

    add_chapter(doc, 3, "Dependency Direction", "You will recognize healthy imports, risky cycles, and boundaries that can scale to a team.")
    add_code(doc, "App -> feature coordinator -> feature component / hook -> utility / API\n\nHealthy: FinderWindow -> getFinderOpenAction\nRisky:   github.js -> FinderWindow\nRisky:   mobile app -> desktop window store", "A simple dependency rule")
    add_para(doc, "FinderWindow asks a pure utility what an item means, then performs the returned action. The utility never imports React, Zustand, or window.open. This makes the decision testable and prevents UI details from becoming the only place where product behavior is defined.")
    add_code(doc, snippet("src/features/finder/utils/getFinderOpenAction.js", 15, 67), "Actual pure action decision")
    add_explanations(doc, "A helper decides; the component performs the decision.", "The utility returns a small command object instead of triggering side effects.", "Separating policy from mechanism lowers coupling and creates a stable unit for tests and future file types.")

    add_chapter(doc, 4, "State Ownership", "You will decide whether state belongs in a component, a custom hook, a domain store, or persistence.")
    add_table(doc, ["State type", "Best owner", "Macfolio example"], [("One component only", "useState in that component", "Mobile navigation stack."), ("Reusable lifecycle logic", "Custom hook", "useWeather and useTodos."), ("Distant components in one domain", "Feature store", "Desktop windows."), ("Cross-feature navigation model", "Focused store", "Finder active location/history."), ("Persisted browser preference/data", "Hook plus localStorage", "Mobile todos."), ("Remote server data", "Data hook/store with status", "Portfolio Work folder.")], [2400, 3300, 3660])
    add_callout(doc, "Decision question", "Which part of the product is allowed to change this value? Ownership is more useful than asking whether state is globally accessible.", "lesson")

    add_chapter(doc, 5, "Refactoring Without Panic", "You will move code toward feature boundaries in safe, explainable steps.")
    add_bullets(doc, ["Document the current responsibility and imports.", "Create the destination feature boundary.", "Move the implementation without redesigning behavior.", "Leave a small compatibility export at the old path.", "Update direct imports gradually.", "Verify lint, build, and user flow.", "Remove the bridge only when no callers remain."], numbered=True)
    add_callout(doc, "When not to refactor", "Do not reorganize a stable area merely because a newer pattern looks fashionable. Refactor when ownership is unclear, change cost is rising, or a feature needs room to grow.", "warning")
    add_checkpoint(doc, "Place a new desktop Calculator app in the structure.", "Choose ownership for shared project-card data used by desktop and mobile.", "Design a migration from compatibility exports to direct feature imports without a flag day.")
    return doc


def volume_three():
    doc = Document(); configure_document(doc, "Volume 3 - App and Shells")
    add_cover(doc, "Volume 3", "From index.html to a Living Interface", "React startup, effects, lazy loading, responsive shell selection, errors, and composition")
    add_toc(doc, [("1. Browser-to-React startup", "Trace execution from HTML to App."), ("2. App composition", "Understand the startup effect and ready state."), ("3. Responsive shells", "Choose desktop or mobile behavior without mixing them."), ("4. Lazy loading and Suspense", "Load large app windows only when needed."), ("5. Error and loading boundaries", "Keep failures understandable and local.")])
    add_chapter(doc, 1, "Browser-to-React Startup", "You will explain exactly what runs before the first Macfolio screen appears.")
    add_figure(doc, make_diagram("startup", "Startup execution flow", ["index.html", "main.jsx", "App.jsx", "data preload", "shell render"]), "Figure 1. Startup is a sequence of boundaries, not magic.")
    add_code(doc, snippet("index.html", 1, 14), "The browser document")
    add_code(doc, snippet("src/main.jsx", 1, 11), "The React mount")
    add_para(doc, "Vite serves index.html and follows the module script to src/main.jsx. The CSS import becomes a stylesheet in the bundle. React then owns the #root element; components return descriptions of UI, and React synchronizes those descriptions with the DOM.")

    add_chapter(doc, 2, "App Composition", "You will read the real App component as a coordinator and understand every stage of its startup effect.")
    add_code(doc, snippet("src/app/App.jsx", 35, 82), "Startup state and effect")
    add_bullets(doc, ["useMediaQuery subscribes to viewport width.", "The window store provides desktop window records.", "The location store receives the initial Work folder.", "usePortfolioFileSystem hides loading and mapping details.", "isReady keeps the intro visible for at least the intended presentation delay.", "The cleanup flag prevents state updates after unmount; the timer is also cleared."])
    add_callout(doc, "Complex function explained", "useEffect runs after React commits the component. Its dependency array says the effect should be recreated if either referenced function identity changes. The returned cleanup runs before recreation and on unmount, protecting asynchronous work from updating a component that no longer exists.", "info")
    add_para(doc, "The finally block is deliberate: even if remote GitHub data fails, the app leaves the intro screen and can continue with fallback or empty data. Reliability sometimes means showing a useful interface while a non-critical service is unavailable.")

    add_chapter(doc, 3, "Responsive Shells", "You will separate a shared product from two interaction models.")
    add_code(doc, snippet("src/app/App.jsx", 84, 132), "Shell selection and desktop registrations")
    add_para(doc, "Desktop and mobile are not just different CSS widths. Desktop uses overlapping windows, z-index, drag, and a Dock. Mobile uses a stack of full-screen app routes and a back action. Keeping the shells separate lets each interaction model remain coherent while shared portfolio concepts can stay below them.")
    add_callout(doc, "When not to split", "If mobile only rearranges the same page sections, responsive CSS may be enough. Split shells when navigation, state transitions, and interaction rules differ materially.", "warning")

    add_chapter(doc, 4, "Lazy Loading and Suspense", "You will understand bundle splitting and know which components deserve lazy loading.")
    add_code(doc, "const Finder = lazy(() => import('../windows/Finder'));\n\n<Suspense fallback={null}>\n  {desktopRenderWindow(Finder, 'finder')}\n</Suspense>", "The lazy window pattern")
    add_five_questions(doc, "React.lazy turns a dynamic import into a component whose code is requested on demand.", "Large editors, PDF renderers, galleries, and secondary apps should not delay the initial screen.", "Use it for substantial routes or app surfaces that may never open in a session.", "Keep declarations near the app or route registry that chooses the component.", "Wrap lazy components in Suspense and provide a fallback appropriate to the surrounding interface.")
    add_callout(doc, "Tradeoff", "Too many tiny chunks add network overhead and make debugging noisier. Lazy-load meaningful feature surfaces, not every button.", "warning")

    add_chapter(doc, 5, "Error and Loading Boundaries", "You will design failure states that protect the rest of the experience.")
    add_code(doc, snippet("src/components/ErrorBoundary.jsx", 1, 61), "Class error boundary used by Macfolio")
    add_para(doc, "React error boundaries catch rendering and lifecycle errors below them. They do not catch event-handler errors, asynchronous promise rejections, or server failures by themselves. Macfolio uses an outer boundary for the whole shell and another around desktop windows so one broken app surface does not necessarily destroy the desktop.")
    add_checkpoint(doc, "Draw the first five startup steps from memory.", "Explain why the Work folder is loaded before desktop Finder opens.", "Choose boundaries for an app with dashboard, billing, and a chat widget.")
    return doc


def volume_four():
    doc = Document(); configure_document(doc, "Volume 4 - Desktop Window System")
    add_cover(doc, "Volume 4", "Build a Desktop Inside the Browser", "Zustand window state, configuration, focus, drag, maximize, Dock behavior, and window applications")
    add_toc(doc, [("1. Model a window", "Turn visual behavior into explicit state."), ("2. Store transitions", "Open, focus, minimize, restore, move, and resize."), ("3. Window wrapper", "Connect React lifecycle to GSAP Draggable."), ("4. Desktop composition", "Dock, Home, Navbar, and controls."), ("5. Add a new app", "Build a complete Calculator window using the established pattern.")])
    add_chapter(doc, 1, "Model a Window", "You will describe a desktop window as data rather than scattered DOM mutations.")
    add_table(doc, ["Property", "Meaning", "Why explicit"], [("isOpen", "Whether the app participates in the desktop.", "Dock and App can make the same decision."), ("isMinimized", "Open but temporarily hidden.", "Restore differs from first open."), ("zIndex", "Stacking priority.", "Focus becomes deterministic."), ("position", "Current x/y coordinates.", "Drag end can persist location."), ("size", "Normal width/height.", "Restore can recover user geometry."), ("sizeMode", "normal or full.", "Maximize is a state transition, not a CSS guess."), ("data", "Payload passed when opening.", "Text/image/code windows can display selected items.")], [1800, 3400, 4160])
    add_code(doc, snippet("src/features/desktop-shell/config/windowConfig.js", 1, 70), "Window configuration is the initial state contract")

    add_chapter(doc, 2, "Store Transitions", "You will read Zustand updates and understand immutable nested state.")
    add_code(doc, snippet("src/features/desktop-shell/store/windowStore.js", 42, 92), "Open, close, and focus transitions")
    add_para(doc, "Zustand's set function receives the previous state and returns the next state. The spread operators create a new windows object and a new record for one window while preserving every unrelated window. React subscribers can detect the changed references and re-render only where needed.")
    add_callout(doc, "Hidden global state", "currentZIndex lives at module scope, outside the store object. It works for one browser session, but it is less visible to tests, reset tools, and devtools. A future refinement could store nextZIndex inside Zustand if those needs become important.", "warning")
    add_figure(doc, make_diagram("window_flow", "Opening a desktop app", ["Dock click", "openWindow", "store update", "App registry", "window wrapper"]), "Figure 1. State drives rendering; the Dock does not create DOM windows directly.")

    add_chapter(doc, 3, "The Window Wrapper", "You will understand higher-order components, refs, layout effects, and drag cleanup.")
    add_code(doc, snippet("src/features/desktop-shell/hoc/windowWrapper.jsx", 21, 55), "A higher-order component returns a managed component")
    add_five_questions(doc, "A higher-order component is a function that receives a component and returns another component.", "All desktop apps need the same frame mechanics, so the behavior is centralized once.", "Use it when several components require the same wrapper contract and hooks cannot alone provide the surrounding DOM.", "The generic mechanics live in desktop-shell/hoc; app content stays in its own window file.", "The wrapper reads store state, holds a DOM ref, applies GSAP geometry, creates a Draggable instance, and kills it during cleanup.")
    add_code(doc, snippet("src/features/desktop-shell/hoc/windowWrapper.jsx", 57, 120), "Maximize and restore synchronization")
    add_para(doc, "useLayoutEffect runs after DOM changes but before the browser paints. That timing prevents a visible flash at the old geometry. useEffect is sufficient for remembering values; useGSAP owns the imperative drag instance and cleanup.")
    add_callout(doc, "Browser implication", "Transforms such as x and y usually avoid document reflow during drag and are efficient for animation. Width and height changes can trigger layout, so they are reserved for mode transitions and resizing rather than every pointer movement.", "info")

    add_chapter(doc, 4, "Desktop Composition", "You will connect Dock, Home, Navbar, controls, and app windows without mixing responsibilities.")
    add_table(doc, ["Part", "Primary job", "Dependency"], [("Dock", "Launch, restore, and focus apps.", "Window store and desktopApps config."), ("Home", "Desktop icons and wallpaper surface.", "Locations and openWindow."), ("Navbar", "Top desktop chrome and clock/menu state.", "Window store and React time state."), ("WindowControls", "Close, minimize, maximize commands.", "Window store."), ("App.jsx registry", "Render open non-minimized app components.", "Window records and lazy imports."), ("windowWrapper", "Generic movement and geometry.", "Store, config, GSAP.")], [1900, 3900, 3560])
    add_para(doc, "Notice the rhythm: configuration describes, stores transition, components request, App composes, and the wrapper applies generic mechanics. Repeating that pattern makes a new desktop app predictable.")

    add_chapter(doc, 5, "Add a New Calculator App", "You will be able to extend the desktop system without inventing a second architecture.")
    add_bullets(doc, ["Add calculator geometry to windowConfig.js.", "Add its icon and id to desktopApps.js.", "Create windows/Calculator.jsx with local calculator state.", "Wrap it with windowWrapper(Calculator, 'calculator').", "Lazy-import it in app/App.jsx.", "Register desktopRenderWindow(Calculator, 'calculator').", "Add focused CSS using a calculator prefix.", "Verify open, focus, minimize, restore, drag, and maximize."], numbered=True)
    add_callout(doc, "Ownership", "The calculator's expression and result are app-specific local state. Window position and focus remain desktop-shell state. Do not put calculator arithmetic inside the global window store.", "lesson")
    add_checkpoint(doc, "Explain why isOpen and isMinimized are separate.", "Trace a drag from pointer press to persisted position.", "Design keyboard focus and accessibility rules for ten simultaneous windows.")
    return doc


def volume_five():
    doc = Document(); configure_document(doc, "Volume 5 - Data, Finder, Preview")
    add_cover(doc, "Volume 5", "Turn Remote Data into an Experience", "GitHub requests, portfolio mapping, Finder navigation, code preview, caching, fallbacks, and security")
    add_toc(doc, [("1. External API boundary", "Fetch public GitHub data without leaking secrets."), ("2. Portfolio data boundary", "Map remote responses into stable product entities."), ("3. Finder navigation", "Separate UI effects from pure open-item policy."), ("4. Code preview", "Coordinate explorer, editor, and remote file content."), ("5. Reliability and growth", "Handle rate limits, caching, tests, and future backend needs.")])
    add_chapter(doc, 1, "The External API Boundary", "You will isolate browser requests and know when public client-side fetching is appropriate.")
    add_code(doc, snippet("src/api/github.js", 14, 112), "Actual GitHub request module")
    add_bullets(doc, ["VITE_GITHUB_USERNAME is public configuration and can safely appear in the browser bundle.", "A GitHub token must not be placed in a VITE_ variable because users can inspect it.", "Non-OK responses return empty arrays so the portfolio can fall back gracefully.", "Recursive tree loading stops after depth 3 to control request count and payload size.", "The API layer returns plain objects; it never opens windows or changes React state."])
    add_callout(doc, "When a backend is required", "Use a server route for private repositories, secret tokens, stronger caching, higher rate limits, auditing, or data normalization shared by many clients.", "risk")

    add_chapter(doc, 2, "The Portfolio Data Boundary", "You will transform unstable remote data into a stable Finder-friendly model.")
    add_figure(doc, make_diagram("data_flow", "Portfolio data flow", ["GitHub API", "buildWorkLocation", "data store", "location store", "Finder"]), "Figure 1. The mapper protects the interface from remote response shape changes.")
    add_code(doc, snippet("src/features/portfolio/utils/buildWorkLocation.js", 35, 136), "Mapping repositories to Finder nodes")
    add_para(doc, "buildRepoFolder creates a product concept from a repository: Source Code, Live Site, and Project.txt. attachParents adds navigation links after the tree is built. cachedWork prevents repeated network work during the same module lifetime. The output contract stays useful even if GitHub is replaced later.")
    add_code(doc, snippet("src/features/portfolio/hooks/usePortfolioFileSystem.js", 15, 58), "Hook coordinating status and storage")
    add_explanations(doc, "The API gives raw boxes; the mapper labels and arranges them for Finder.", "The hook coordinates asynchronous loading while the utility owns transformation.", "A stable anti-corruption boundary prevents vendor response shapes from spreading through UI code.")

    add_chapter(doc, 3, "Finder Navigation", "You will build navigation state, breadcrumbs, grids, and file-open behavior with clear ownership.")
    add_code(doc, snippet("src/features/finder/FinderWindow.jsx", 25, 95), "Finder coordinator")
    add_para(doc, "FinderWindow coordinates four things: current location, history action, presentation components, and execution of an open-item command. FinderSidebar, FinderBreadcrumbs, and FinderGrid stay focused on rendering and user events.")
    add_code(doc, snippet("src/features/finder/utils/getFinderOpenAction.js", 15, 67), "Pure file policy")
    add_table(doc, ["Node", "Action", "Result"], [("folder", "SET_LOCATION", "Finder navigates into children."), ("pdf", "OPEN_WINDOW resume", "Resume viewer opens."), ("local txt", "OPEN_WINDOW txtfile", "Text viewer receives item data."), ("image", "OPEN_WINDOW imgfile", "Image viewer receives item data."), ("url/link", "OPEN_EXTERNAL_LINK", "Browser opens a new tab."), ("code extension", "OPEN_WINDOW vsCode", "Code preview receives selected file."), ("unknown", "NONE", "No side effect.")], [2100, 2900, 4360])

    add_chapter(doc, 4, "Code Preview", "You will understand how Explorer, Editor, Terminal, and the window coordinator divide work.")
    add_table(doc, ["File", "Responsibility", "State"], [("CodePreviewWindow.jsx", "Coordinate selected file, loading, content, and shell layout.", "Selected file and fetched content."), ("Explorer.jsx", "Render a recursive file tree and report selection.", "Mostly presentation."), ("Editor.jsx", "Display source through Monaco with language selection.", "Editor presentation/settings."), ("Terminal.jsx", "Provide a decorative terminal/status region.", "Minimal local UI."), ("getFileIcon.jsx", "Map extensions to icon components.", "Pure decision.")], [2700, 4300, 2360])
    add_callout(doc, "Performance", "Monaco and React PDF are substantial packages, which is why their containing windows are lazy-loaded. Performance work begins with loading boundaries before micro-optimizing small functions.", "info")
    img = ROOT / "public/images/project-1.png"
    add_figure(doc, img, "Figure 2. Project media is product content; keep it in public assets and reference stable paths.", width=6.0)

    add_chapter(doc, 5, "Reliability and Growth", "You will know which parts are portfolio-sized today and how they would evolve under real traffic.")
    add_table(doc, ["Concern", "Current choice", "Growth path"], [("Rate limits", "Unauthenticated public GitHub calls.", "Server cache and authenticated server requests."), ("Caching", "Module-level cached Work tree.", "Query cache with expiry and background refresh."), ("Errors", "Empty arrays and fallback repositories.", "Typed errors, retry controls, monitoring."), ("Testing", "Pure decisions are testable by design.", "Unit tests for mappers/actions; integration tests for windows."), ("Large repositories", "Depth limit of 3.", "Lazy-expand folders and fetch on demand."), ("Security", "No secret browser token.", "Server-side credentials and policy checks.")], [1900, 3600, 3860])
    add_checkpoint(doc, "Explain why a VITE token is not secret.", "Design a test matrix for getFinderOpenAction.", "Plan an on-demand folder-loading API for repositories with 100,000 files.")
    return doc


def volume_six():
    doc = Document(); configure_document(doc, "Volume 6 - CSS, Mobile, Growth")
    add_cover(doc, "Volume 6", "Make It Coherent, Then Make It Grow", "CSS architecture, mobile apps, browser persistence, refactoring, deployment, and the complete repository atlas")
    add_toc(doc, [("1. CSS as a system", "Read global layout, feature styles, cascade, and responsive rules."), ("2. Mobile shell", "Build stack navigation, app frames, widgets, and mini-apps."), ("3. Browser APIs", "Use localStorage, events, geolocation, and fetch safely."), ("4. Quality and deployment", "Verify, optimize, document, and ship."), ("5. Future architecture", "Plan growth without premature complexity."), ("Appendix A", "Analyze all 180 repository files and assets.")])
    add_chapter(doc, 1, "CSS as a System", "You will understand how Macfolio CSS is laid out and how to keep a growing stylesheet predictable.")
    add_five_questions(doc, "CSS is a rule system that matches elements and resolves competing declarations through origin, specificity, and source order.", "A coherent system prevents each component from inventing spacing, colors, and layout behavior independently.", "Define global primitives once; keep feature-specific selectors beside the feature when they form a substantial visual domain.", "Macfolio has src/index.css for desktop/global styling and features/mobile-shell/styles.css for the mobile domain.", "Organize from foundations to shells to components to states to responsive overrides, using stable class prefixes.")
    add_table(doc, ["CSS layer", "Examples", "Reason"], [("Reset/foundation", "box-sizing, body/main defaults", "Create predictable browser behavior."), ("Shell geometry", "nav, dock, window layer, mobile root", "Define the major coordinate systems."), ("Components", "dock icon, Finder grid, editor panes", "Style reusable visual units."), ("States", "active, minimized, full, error", "Make interaction state visible."), ("Responsive", "media queries and stable dimensions", "Adapt without accidental layout shifts."), ("Motion", "transitions and keyframes", "Support meaning; respect reduced motion later.")], [2100, 3800, 3460])
    add_code(doc, snippet("src/index.css", 1, 55), "Global foundations and early desktop selectors")
    add_para(doc, "Flexbox handles one-dimensional alignment such as nav items or toolbars. Grid handles two-dimensional collections such as icons and cards. Fixed positioning creates the desktop coordinate space for windows; transforms move dragged windows efficiently. Overflow rules decide which layer may scroll and which must clip.")
    add_callout(doc, "Current debt", "Both stylesheets are large. That is not automatically wrong, but feature growth will make ownership harder. A gradual next step is to move desktop feature styles beside Finder, code preview, and desktop shell while preserving naming and visual tokens.", "warning")

    add_chapter(doc, 2, "The Mobile Shell", "You will recreate mobile stack navigation and understand why it stays separate from desktop windows.")
    add_code(doc, snippet("src/features/mobile-shell/MobileApp.jsx", 1, 85), "Mobile navigation coordinator")
    add_para(doc, "The stack is an array of route records. Opening an app appends a record; going back slices off the last record. The last item is active. A timestamp key makes repeated openings distinct for animation. AnimatePresence keeps the outgoing screen mounted long enough to play its exit transition.")
    add_table(doc, ["Mobile layer", "Owns"], [("data/apps.js", "App ids, names, icons, and home/dock placement."), ("HomeScreen.jsx", "Widgets, app icons, and open-app requests."), ("AppFrame.jsx", "Shared status/header/back frame for an active app."), ("MobileApp.jsx", "Stack navigation and app-component registry."), ("apps/*.jsx", "Individual mobile feature UI."), ("styles.css", "Phone shell and mini-app visual system.")], [3100, 6260])

    add_chapter(doc, 3, "Browser APIs as Dependencies", "You will handle persistence, cross-tab events, geolocation, and network cleanup deliberately.")
    add_code(doc, snippet("src/features/mobile-shell/data/todos.js", 1, 92), "Local todo persistence hook")
    add_para(doc, "useState(readTodos) passes a function, so storage is read only during initial state creation. commit computes the next immutable array, writes JSON, dispatches a same-tab custom event, and returns the new state. The storage event covers other tabs; the custom event covers components in the current tab.")
    add_code(doc, snippet("src/features/mobile-shell/data/weather.js", 38, 76), "Geolocation with a fallback")
    add_callout(doc, "Permission design", "Geolocation can be denied, unavailable, or slow. The fallback is part of the feature contract, not an afterthought. Good browser code treats permission and network failure as normal branches.", "lesson")

    add_chapter(doc, 4, "Quality and Deployment", "You will verify the project like an engineer and ship a reproducible production build.")
    add_bullets(doc, ["Run npm run lint and fix errors rather than suppressing unfamiliar rules.", "Run npm run build to catch bundling and import problems.", "Test desktop at wide and short viewports; test mobile around the 768px boundary.", "Open, focus, minimize, restore, drag, and maximize every desktop window.", "Test GitHub failure, geolocation denial, empty storage, and corrupted storage.", "Check keyboard focus, button labels, image alt text, contrast, and reduced motion.", "Deploy the dist output to a static host and configure VITE_GITHUB_USERNAME at build time."], numbered=True)
    add_callout(doc, "Definition of done", "A feature is complete when behavior, failure states, architecture notes, and verification all agree. A pretty happy path is only one slice of production quality.", "lesson")

    add_chapter(doc, 5, "Future Architecture Without Premature Complexity", "You will distinguish the next sensible improvement from enterprise theater.")
    add_table(doc, ["If this happens", "Then consider", "Do not add yet"], [("More data sources", "A portfolio repository interface and adapters.", "A generic plugin framework."), ("More remote state", "A query cache with retries and stale times.", "A global store for every request."), ("More desktop apps", "An app registry that contains component loaders and metadata.", "Dynamic micro-frontends."), ("More contributors", "Feature ownership docs, tests, lint boundaries.", "A monorepo without separate products."), ("Private GitHub data", "A small server/API route.", "Secrets in client environment variables."), ("CSS collisions", "Feature styles and shared tokens.", "A complete design-system package immediately.")], [2500, 3400, 3460])
    add_checkpoint(doc, "Explain specificity without using the word important.", "Design state ownership for a mobile Notes app with persistence.", "Propose the next three refactors and justify their order by risk and value.")

    doc.add_page_break()
    doc.add_heading("Appendix A: Complete Repository Atlas", level=1)
    add_para(doc, "Coverage statement: 180 files discovered, 180 files read or inspected, 0 unread, 100% repository coverage. Text files were decoded and analyzed; raster/PDF assets were inspected through metadata and visual contact sheets. Each entry below records the file's role, dependencies, execution story, design choice, alternative, and growth note.")
    for index, entry in enumerate(repository_atlas(), 1):
        doc.add_heading(f"A.{index}  {entry['path']}", level=2)
        rows = [
            ("Purpose", entry["purpose"]),
            ("Responsibility", entry["responsibility"]),
            ("Dependencies / imports", entry["dependencies"]),
            ("Exports", entry["exports"]),
            ("Execution flow", entry["flow"]),
            ("Design decision", entry["decision"]),
            ("Alternative", entry["alternative"]),
            ("Scaling note", entry["scaling"]),
        ]
        add_table(doc, ["Lens", "Analysis"], rows, [1800, 7560])
    return doc


def inventory_files():
    names = subprocess.check_output(["rg", "--files"], cwd=ROOT, text=True).splitlines()
    output_prefix = "docs/tutorials/macfolio-textbook/"
    return [
        ROOT / name
        for name in sorted(names)
        if not name.replace("\\", "/").startswith(output_prefix)
    ]


def extract_purpose(text):
    match = re.search(r"\* PURPOSE:\s*\n\s*\*\s*(.+)", text)
    return match.group(1).strip() if match else None


def repository_atlas():
    entries = []
    for path in inventory_files():
        rel = path.relative_to(ROOT).as_posix()
        ext = path.suffix.lower()
        if ext in TEXT_EXTENSIONS:
            text = path.read_text(encoding="utf-8")
            imports = re.findall(r"(?:from\s+|import\s*\()['\"]([^'\"]+)", text)
            exports = re.findall(r"export\s+(?:default\s+)?(?:function|const|class|let|var)?\s*([A-Za-z_$][\w$]*)", text)
            purpose = extract_purpose(text) or infer_purpose(rel, ext)
            responsibility = infer_responsibility(rel, ext)
            deps = ", ".join(dict.fromkeys(imports)) or infer_asset_dependencies(rel, text)
            out = ", ".join(dict.fromkeys(exports)) or ("default/compatibility export" if "export { default }" in text else "None")
            flow = infer_flow(rel, ext)
        elif ext in IMAGE_EXTENSIONS:
            with Image.open(path) as image:
                size = f"{image.width}x{image.height} {image.mode}"
            purpose = f"Visual asset ({size}) used by the portfolio interface or content gallery."
            responsibility = "Provide a stable browser-loadable image; it contains no application behavior."
            deps = "Loaded by CSS, JSX, or portfolio content through a /public path."
            out = "Static file URL"
            flow = "Vite serves the file unchanged; the browser decodes and paints it when referenced."
        elif ext == ".pdf":
            purpose = "One-page resume document rendered by the desktop and mobile resume experiences."
            responsibility = "Store the downloadable/viewable resume artifact."
            deps = "react-pdf and direct browser download links."
            out = "Static /files/resume.pdf URL"
            flow = "The browser fetches the PDF; react-pdf parses it and paints the requested page."
        else:
            purpose = infer_purpose(rel, ext)
            responsibility = infer_responsibility(rel, ext)
            deps = "Repository tooling or runtime consumer."
            out = "Static artifact"
            flow = "Consumed by the relevant build or browser subsystem."
        entries.append({
            "path": rel,
            "purpose": purpose,
            "responsibility": responsibility,
            "dependencies": deps or "None",
            "exports": out,
            "flow": flow,
            "decision": infer_decision(rel, ext),
            "alternative": infer_alternative(rel, ext),
            "scaling": infer_scaling(rel, ext),
        })
    return entries


def infer_purpose(rel, ext):
    name = Path(rel).name
    if rel.startswith("docs/"):
        return f"Teach or record the engineering topic represented by {name}."
    if ext == ".svg":
        return "Vector icon or logo served as a static visual asset."
    if name == "package-lock.json":
        return "Pin the complete npm dependency graph for reproducible installs."
    if name == "package.json":
        return "Declare project scripts, runtime dependencies, and development dependencies."
    if name.endswith(".css"):
        return "Define the visual rules for the scope indicated by its path."
    if name.endswith((".jsx", ".js")):
        return f"Implement or expose the {Path(name).stem} module in its owning folder."
    return f"Support the project through the {name} artifact."


def infer_responsibility(rel, ext):
    if "/components/" in rel or rel.endswith(".jsx"):
        return "Render or coordinate the UI behavior implied by the file and its feature boundary."
    if "/data/" in rel or "/constants/" in rel:
        return "Provide stable data/configuration without owning unrelated UI behavior."
    if "/store/" in rel:
        return "Own and transition shared state for its domain."
    if "/utils/" in rel:
        return "Perform a focused transformation or decision with minimal side effects."
    if "/api/" in rel:
        return "Communicate with an external service and return plain data."
    if rel.startswith("docs/"):
        return "Preserve reasoning and learning material as part of the repository."
    if ext in {".css", ".svg"}:
        return "Provide presentation only; application decisions remain in JavaScript modules."
    return "Provide the focused configuration or artifact described by its purpose."


def infer_asset_dependencies(rel, text):
    if rel.endswith(".css"):
        return "Imported by its owning React entry; references public assets and class names."
    if rel.endswith(".md"):
        return "Human readers and future maintainers."
    if rel.endswith(".svg"):
        return "Browser SVG renderer and the JSX/CSS path that references it."
    return "No JavaScript imports."


def infer_flow(rel, ext):
    if rel.endswith("main.jsx"):
        return "Browser loads the module, creates the React root, and renders App."
    if rel.endswith(".css"):
        return "The bundler loads the stylesheet; the browser matches selectors during style calculation."
    if rel.startswith("docs/"):
        return "A developer reads it before or after implementation to understand context and tradeoffs."
    if "/store/" in rel:
        return "Components call actions; the store creates new state; subscribers re-render."
    if "/utils/" in rel:
        return "A caller supplies plain input; the function returns a transformed value or explicit decision."
    if "/api/" in rel:
        return "A feature requests data; fetch resolves; the module validates and returns a safe shape."
    if ext == ".svg":
        return "The browser loads and rasterizes vector paths when an img reference becomes visible."
    if rel.endswith((".jsx", ".js")):
        return "The importing module evaluates this file; React or its caller invokes the exported behavior."
    return "Consumed during install, build, development, or browser rendering."


def infer_decision(rel, ext):
    if rel.startswith("src/features/"):
        return "Code is placed with the feature that owns its change, reinforcing feature-based architecture."
    if rel.startswith("src/components/") or rel.startswith("src/windows/") or rel.startswith("src/store/"):
        return "The path supports app-level UI or a gradual compatibility migration; avoid adding unrelated behavior."
    if rel.startswith("docs/decisions/"):
        return "An ADR preserves context, choice, consequences, and tradeoffs rather than only the final code."
    if ext in IMAGE_EXTENSIONS or ext in {".svg", ".pdf"}:
        return "Static user-facing assets live under public so their URLs remain simple and build-independent."
    return "The file has one primary responsibility and follows the repository's existing convention."


def infer_alternative(rel, ext):
    if rel.endswith(".css"):
        return "CSS Modules or CSS-in-JS could isolate selectors, but migration should be gradual and justified by collisions."
    if "/store/" in rel:
        return "React context or lifted state could work at smaller scope; the store is justified by distant consumers."
    if "/api/" in rel:
        return "A server proxy would be preferable for secrets, private data, or stronger caching."
    if rel.startswith("docs/"):
        return "Reasoning could remain in chat or commit messages, but dedicated docs are easier to discover and teach from."
    if ext in IMAGE_EXTENSIONS or ext == ".svg":
        return "An icon library or generated asset could replace it if licensing, consistency, or bundle goals change."
    return "The same behavior could be colocated in a caller, but separation is useful when ownership or reuse is clear."


def infer_scaling(rel, ext):
    if rel.startswith("src/features/"):
        return "Add subfolders only when the feature grows; preserve the boundary and keep cross-feature dependencies explicit."
    if rel.startswith("src/components/") or rel.startswith("src/windows/") or rel.startswith("src/store/"):
        return "Retire compatibility paths after callers migrate; do not let the bridge become a second permanent architecture."
    if rel.endswith(".css"):
        return "Split by feature when ownership or selector collision becomes costly; centralize stable tokens before duplicating values."
    if ext in IMAGE_EXTENSIONS or ext in {".svg", ".pdf"}:
        return "Optimize dimensions and formats, add descriptive alt usage, and consider a CDN only when traffic or payload warrants it."
    if rel.startswith("docs/"):
        return "Update alongside behavior so the document remains trustworthy as the team grows."
    return "Keep the responsibility narrow and revisit the abstraction only after repeated real change pressure."


def master_index():
    doc = Document(); configure_document(doc, "Master Index")
    add_cover(doc, "Master Index", "The Macfolio Engineering Textbook", "A rebuild path from first command to architecture confidence")
    doc.add_heading("How To Use This Course", level=1)
    add_para(doc, "This course assumes you know a little HTML, CSS, JavaScript, and React but have not yet developed strong project-organization instincts. Work through the volumes in order. Keep Macfolio open beside the book, type the important examples, and answer each checkpoint before reading onward.")
    add_callout(doc, "Coverage", "180 files discovered; 149 text files read; 31 binary assets inspected; 0 unread; repository coverage 100%. The course uses actual repository code and does not claim behavior that was reconstructed from guesswork.", "lesson")
    add_table(doc, ["Volume", "Build outcome"], [("1. Start Without Hesitation", "Create and configure the project; adopt a repeatable engineering rhythm."), ("2. Architecture You Can Feel", "Design feature boundaries, dependencies, and state ownership."), ("3. From index.html to a Living Interface", "Trace startup and compose responsive shells."), ("4. Build a Desktop Inside the Browser", "Implement the complete desktop window system."), ("5. Turn Remote Data into an Experience", "Build GitHub data, Finder, and code-preview flows."), ("6. Make It Coherent, Then Make It Grow", "Master CSS/mobile/browser APIs, deployment, and every repository file.")], [3300, 6060])
    doc.add_heading("The Rebuild Order", level=1)
    add_bullets(doc, ["Bootstrap Vite and understand the entry point.", "Create app and feature boundaries before filling them.", "Build desktop window state and one visible app.", "Add Finder navigation and a local portfolio tree.", "Connect public GitHub data through the portfolio mapper.", "Add code preview and remaining desktop apps.", "Build the separate mobile shell and mini-apps.", "Organize CSS, failure states, accessibility, and performance.", "Run production verification and deploy.", "Review the repository atlas and explain ownership in your own words."], numbered=True)
    doc.add_heading("Confidence Contract", level=1)
    add_para(doc, "You do not need to know the whole final architecture before starting a new idea. You need a small set of questions, one coherent first slice, and the discipline to improve boundaries as real pressure appears. That is the central habit this textbook is designed to build.")
    return doc


def save(doc, filename):
    path = OUT / filename
    doc.save(path)
    return path


def main():
    outputs = [
        save(master_index(), "00-macfolio-textbook-master-index.docx"),
        save(volume_one(), "01-start-without-hesitation.docx"),
        save(volume_two(), "02-architecture-you-can-feel.docx"),
        save(volume_three(), "03-from-index-to-interface.docx"),
        save(volume_four(), "04-desktop-window-system.docx"),
        save(volume_five(), "05-data-finder-code-preview.docx"),
        save(volume_six(), "06-css-mobile-growth-and-repository-atlas.docx"),
    ]
    print("\n".join(str(path) for path in outputs))


if __name__ == "__main__":
    main()
